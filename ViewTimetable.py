
import sys
import sqlite3
from PyQt5.QtWidgets import (QWidget, QVBoxLayout, QTableWidget, QInputDialog,
                             QTableWidgetItem, QPushButton, QLabel,
                             QMessageBox, QApplication, QFormLayout, 
                             QLineEdit, QHeaderView, QComboBox, QHBoxLayout, QFileDialog,QTimeEdit)
from PyQt5.QtCore import pyqtSignal, Qt, QRect,QTime
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from PyQt5.QtGui import QPainter, QFont,QPen
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from PyQt5.QtWidgets import QMessageBox, QFileDialog
from reportlab.lib.pagesizes import landscape
import textwrap
from reportlab.lib.pagesizes import A4, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
from reportlab.lib.utils import simpleSplit
from fpdf import FPDF
from reportlab.lib.pagesizes import A4, landscape
from fpdf import FPDF
from docx import Document
import os
from docx.shared import Pt



def fetch_query_results(query, params=()):
    connection = sqlite3.connect('timetable.db')
    cursor = connection.cursor()
    
    cursor.execute(query, params)
    results = cursor.fetchall()
    
    connection.close()
    return results

def delete_record(record_id, table_name):
    connection = sqlite3.connect('timetable.db')
    cursor = connection.cursor()
    
    query = f"DELETE FROM {table_name} WHERE ID = ?"
    cursor.execute(query, (record_id,))
    
    connection.commit()
    connection.close()

class ViewTimetableWindow(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("View Timetable")
        self.setGeometry(200, 200, 800, 600)
        self.layout = QVBoxLayout(self)
        self.showMaximized()

        # Filters
        self.filter_layout = QHBoxLayout()
        self.department_filter = QComboBox(self)
        self.department_filter.addItem("All Departments")
        self.semester_filter = QComboBox(self)
        self.semester_filter.addItem("All Semesters")
        self.teacher_filter = QComboBox(self)
        self.teacher_filter.addItem("All Teachers")
        self.session_filter = QComboBox(self)  # Session filter
        self.session_filter.addItem("All Sessions")

        self.filter_layout.addWidget(QLabel("Department:", self))
        self.filter_layout.addWidget(self.department_filter)
        self.filter_layout.addWidget(QLabel("Semester:", self))
        self.filter_layout.addWidget(self.semester_filter)
        self.filter_layout.addWidget(QLabel("Teacher:", self))
        self.filter_layout.addWidget(self.teacher_filter)
        self.filter_layout.addWidget(QLabel("Session:", self))  # Add session label
        self.filter_layout.addWidget(self.session_filter)        # Add session filter ComboBox
        
        # Connect filter change to filtering function
        self.department_filter.currentIndexChanged.connect(self.apply_filters)
        self.semester_filter.currentIndexChanged.connect(self.apply_filters)
        self.teacher_filter.currentIndexChanged.connect(self.apply_filters)
        self.session_filter.currentIndexChanged.connect(self.apply_filters)  # Connect session filter

        self.layout.addLayout(self.filter_layout)
        ##
        # Print Button
        self.print_button = QPushButton("Print To PDF", self)
        self.print_button.clicked.connect(self.generate_timetable_report)
        self.layout.addWidget(self.print_button)

        # Debug label
        self.debug_label = QLabel("Loading timetable data...", self)
        self.layout.addWidget(self.debug_label)

        # Table widget
        self.table_widget = QTableWidget(self)
        self.layout.addWidget(self.table_widget)
        self.table_widget.horizontalHeader().setStretchLastSection(True)
        self.table_widget.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)

        # Load data and filters
        self.load_filter_options()
        self.load_timetable_data()

    def load_filter_options(self):
        """Load unique filter options from the database into combo boxes."""
        try:
            # Populate Department filter
            departments = fetch_query_results("SELECT DISTINCT Department FROM Timetable")
            for department in departments:
                self.department_filter.addItem(department[0])

            # Populate Semester filter
            semesters = fetch_query_results("SELECT DISTINCT Semester FROM Timetable")
            for semester in semesters:
                self.semester_filter.addItem(semester[0])

            # Populate Teacher filter
            teachers = fetch_query_results("SELECT DISTINCT Teacher FROM Timetable")
            for teacher in teachers:
                self.teacher_filter.addItem(teacher[0])

            # Populate Session filter
            sessions = fetch_query_results("SELECT DISTINCT Session FROM Timetable")
            for session in sessions:
                self.session_filter.addItem(session[0])

        except Exception as e:
            self.debug_label.setText(f"Error loading filter options: {e}")

    def load_timetable_data(self):
        """Load timetable data from the database and populate the table widget."""
        try:
            query = "SELECT * FROM Timetable"
            results = fetch_query_results(query)

            if results:
                self.debug_label.setText("Timetable data loaded successfully!")
                self.populate_table(results)
            else:
                self.debug_label.setText("No timetable data available.")

        except Exception as e:
            self.debug_label.setText(f"Error loading timetable: {e}")


    def apply_filters(self):
        """Apply the selected filters and update the table display."""
        department = self.department_filter.currentText()
        semester = self.semester_filter.currentText()
        teacher = self.teacher_filter.currentText()
        session = self.session_filter.currentText()

        # Build query with filters
        query = "SELECT * FROM Timetable WHERE 1=1"
        params = []

        if department != "All Departments":
            query += " AND Department = ?"
            params.append(department)
        
        if semester != "All Semesters":
            query += " AND Semester = ?"
            params.append(semester)
        
        if teacher != "All Teachers":
            query += " AND Teacher = ?"
            params.append(teacher)
        
        if session != "All Sessions":
            query += " AND Session = ?"
            params.append(session)

        try:
            results = fetch_query_results(query, params)
            self.populate_table(results)

        except Exception as e:
            self.debug_label.setText(f"Error applying filters: {e}")


    def populate_table(self, results):
        """Populate the table widget with the provided results."""
        self.table_widget.setRowCount(len(results))
        self.table_widget.setColumnCount(len(results[0]) + 2)
        self.table_widget.setHorizontalHeaderLabels(['ID', 'Department', 'Semester', 'Teacher', 
                                                     'Course Title', 'Course Code', 'Classroom', 
                                                     'Start Time', 'End Time', 'Session', 'Update', 'Delete'])

        for row_index, row_data in enumerate(results):
            for col_index, data in enumerate(row_data):
                self.table_widget.setItem(row_index, col_index, QTableWidgetItem(str(data)))

            # Update button
            update_button = QPushButton("Update")
            update_button.clicked.connect(lambda checked, id=row_data[0], data=row_data: self.open_update_window(id, data))
            self.table_widget.setCellWidget(row_index, len(row_data), update_button)

            # Delete button
            delete_button = QPushButton("Delete")
            delete_button.clicked.connect(lambda checked, id=row_data[0]: self.delete_record(id))
            self.table_widget.setCellWidget(row_index, len(row_data) + 1, delete_button)








            from docx import Document
            from PyQt5.QtWidgets import QFileDialog, QMessageBox

        from docx.shared import Pt
        from docx.enum.table import WD_ALIGN_VERTICAL
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        from docx import Document
        import sqlite3  # Assuming you are using SQLite for your database

    def generate_timetable_report(self):
        """Generates the timetable report by applying data to a Word template."""
        
        # Prompt user to save the output
        file_dialog = QFileDialog(self)
        file_path, _ = file_dialog.getSaveFileName(self, "Save Report", "", "Word Files (*.docx)")
        
        if not file_path:
            return  # User canceled the save dialog

        # Load the Word template
        doc = Document("C:/Users/Muhammad Haseeb/Documents/GitHub/College_Time_Table_Software/VBA.docx")

        # Get selected filters
        department = self.department_filter.currentText()
        semester = self.semester_filter.currentText()
        teacher = self.teacher_filter.currentText()
        session = self.session_filter.currentText()

        # Define placeholders and replacements
        replacements = {
            "{{Department}}": department,
            "{{Semester}}": semester,
            "{{Teacher}}": teacher,
            "{{Session}}": session,
        }

        # Replace placeholders in the document
        for paragraph in doc.paragraphs:
            for placeholder, replacement in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, replacement)

        # Fetch timetable data from the database
        connection = sqlite3.connect('timetable.db')  # Update with your database path
        cursor = connection.cursor()

        query = """
        SELECT department, semester, teacher, course_title, course_code, classroom, 
            lecture_start_time, lecture_end_time, session
        FROM Timetable
        WHERE department = ? AND semester = ? AND teacher = ? AND session = ?
    """
        cursor.execute(query, (department, semester, teacher, session))
        timetable_data = cursor.fetchall()

        # Locate the table in the document (assuming the timetable is stored in the first table)
        table = doc.tables[0]
        
        # Table headers (excluding 'ID' column)
        headers = ['Department', 'Semester', 'Teacher', 'Course Title', 'Course Code', 
                'Classroom', 'Start Time', 'End Time', 'Session']
        
        # Fill in timetable data
        for data in timetable_data:
            row = table.add_row().cells
            for i, value in enumerate(data):
                cell = row[i]
                cell.text = str(value)
                cell.paragraphs[0].paragraph_format.word_wrap = True  # Enable word wrap in the cell
                cell.paragraphs[0].font.size = Pt(10)  # Optional: adjust font size
                cell.alignment = WD_ALIGN_VERTICAL.CENTER  # Optional: center-align text vertically

        # Apply auto-adjustment for column width
        for col in table.columns:
            for cell in col.cells:
                cell.width = Pt(100)  # Adjust column width as needed

        # Save the updated document
        doc.save(file_path)
        connection.close()

        QMessageBox.information(self, "Report Saved", "The timetable report was successfully saved.")
























































        




















 




















    def open_update_window(self, record_id, row_data):
        self.update_window = UpdateTimetableWindow(record_id, row_data)  
        self.update_window.update_successful.connect(self.load_timetable_data)
        self.update_window.show()

    def delete_record(self, record_id):
        reply = QMessageBox.question(self, 'Delete Confirmation', f"Are you sure you want to delete record ID {record_id}?",
                                     QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            delete_record(record_id, "Timetable")
            self.load_timetable_data()

class UpdateTimetableWindow(QWidget):
    update_successful = pyqtSignal()

    def __init__(self, record_id, row_data, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Update Timetable Record")
        self.setGeometry(150, 150, 400, 300)

        self.record_id = record_id  
        self.row_data = row_data  
        self.parent = parent  

        self.layout = QFormLayout(self)

        # Department ComboBox populated from the database
        self.department_input = QComboBox(self)
        self.populate_combobox(self.department_input, "Department", row_data[1])
        self.layout.addRow("Department:", self.department_input)

        # Semester ComboBox populated from the database
        self.semester_input = QComboBox(self)
        self.populate_combobox(self.semester_input, "Semester", row_data[2])
        self.layout.addRow("Semester:", self.semester_input)

        # Teacher ComboBox populated from the database
        self.teacher_input = QComboBox(self)
        self.populate_combobox(self.teacher_input, "Teacher", row_data[3])
        self.layout.addRow("Teacher:", self.teacher_input)

        self.course_title_input = QLineEdit(self)
        self.course_title_input.setText(row_data[4])  
        self.layout.addRow("Course Title:", self.course_title_input)

        self.course_code_input = QLineEdit(self)
        self.course_code_input.setText(row_data[5])  
        self.layout.addRow("Course Code:", self.course_code_input)

        self.classroom_input = QLineEdit(self)
        self.classroom_input.setText(row_data[6])  
        self.layout.addRow("Classroom:", self.classroom_input)

        # Use QTimeEdit for start time with 12-hour format
        self.start_time_input = QTimeEdit(self)
        self.start_time_input.setDisplayFormat("hh:mm AP")
        self.start_time_input.setTime(QTime.fromString(row_data[7], "hh:mm AP"))  # Set initial time
        self.layout.addRow("Start Time:", self.start_time_input)

        # Use QTimeEdit for end time with 12-hour format
        self.end_time_input = QTimeEdit(self)
        self.end_time_input.setDisplayFormat("hh:mm AP")
        self.end_time_input.setTime(QTime.fromString(row_data[8], "hh:mm AP"))  # Set initial time
        self.layout.addRow("End Time:", self.end_time_input)

        self.session_input = QLineEdit(self)
        self.session_input.setText(row_data[9])  
        self.layout.addRow("Session:", self.session_input)

        self.update_button = QPushButton("Update", self)
        self.update_button.clicked.connect(self.update_record)
        self.layout.addRow(self.update_button)

        self.setLayout(self.layout)

    def populate_combobox(self, combo_box, column_name, current_value):
        """Populate the combo box with data from the database."""
        query = f"SELECT DISTINCT {column_name} FROM Timetable"
        values = fetch_query_results(query)

        # Add 'All' option for non-specific filtering
        combo_box.addItem(f"All {column_name}s")
        
        for value in values:
            combo_box.addItem(value[0])
        
        # Set the current value from the record if available
        if current_value:
            index = combo_box.findText(current_value)
            if index != -1:
                combo_box.setCurrentIndex(index)

    def update_record(self):
        updated_data = (
            self.department_input.currentText(),
            self.semester_input.currentText(),
            self.teacher_input.currentText(),
            self.course_title_input.text(),
            self.course_code_input.text(),
            self.classroom_input.text(),
            self.start_time_input.time().toString("hh:mm AP"),  # Get selected start time
            self.end_time_input.time().toString("hh:mm AP"),    # Get selected end time
            self.session_input.text()
        )

        update_record("Timetable", self.record_id, updated_data)
        self.update_successful.emit()
        QMessageBox.information(self, "Success", "Record updated successfully!")
        if self.parent:
            self.parent.load_timetable_data()
        self.close()


def update_record(table_name, record_id, updated_data):
    query = f"""UPDATE {table_name} 
                 SET department = ?,           
                     semester = ?, 
                     teacher = ?, 
                     course_title = ?,        
                     course_code = ?, 
                     classroom = ?, 
                     lecture_start_time = ?, 
                     lecture_end_time = ?, 
                     session = ? 
                 WHERE id = ?"""
    
    connection = sqlite3.connect('timetable.db')
    cursor = connection.cursor()

    try:
        cursor.execute(query, (*updated_data, record_id))
        connection.commit()
    except Exception as e:
        print(f"Error updating record: {e}")
    finally:
        connection.close()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = ViewTimetableWindow()
    window.show()
    sys.exit(app.exec_())
